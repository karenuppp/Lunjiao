"""Apply 公文格式 to markdown and produce .docx in sandbox.

Reads /sandbox/input.md, writes /sandbox/output/result.docx.
Fonts: 方正小标宋简体 (title), 黑体 (h1), 楷体_GB2312 (h2), 仿宋_GB2312 (body).
"""
import os
import re
import json
import uuid
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cn_font(run, font_name, size_pt, bold=False):
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def _add_paragraph(doc, text, font_name, size_pt, bold=False,
                   alignment=None, spacing_before=0, spacing_after=0,
                   first_line_indent=False, left_indent_pt=0):
    para = doc.add_paragraph()
    if alignment is not None:
        para.alignment = alignment
    para.paragraph_format.space_before = Pt(spacing_before)
    para.paragraph_format.space_after = Pt(spacing_after)
    para.paragraph_format.line_spacing = Pt(28)
    if left_indent_pt:
        para.paragraph_format.left_indent = Pt(left_indent_pt)
    if first_line_indent:
        para.paragraph_format.first_line_indent = Pt(size_pt * 2)
    clean = _strip_inline_markdown(text)
    run = para.add_run(clean)
    _set_cn_font(run, font_name, size_pt, bold)
    return para


def _strip_inline_markdown(text):
    text = re.sub(r'\!\[.*?\]\(.*?\)', '', text)       # images
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # links
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)        # bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)            # italic
    text = re.sub(r'`{1,3}[^`]*`{1,3}', '', text)      # inline code
    text = re.sub(r'~~(.+?)~~', r'\1', text)            # strikethrough
    return text.strip()


_RE_L1 = re.compile(r'^[一二三四五六七八九十]、')
_RE_L2 = re.compile(r'^（[一二三四五六七八九十]）')
_RE_L3 = re.compile(r'^\d+\.')
_RE_L4 = re.compile(r'^（\d+）')


def _classify_line(line):
    if _RE_L1.match(line):
        return 1
    if _RE_L2.match(line):
        return 2
    if _RE_L3.match(line):
        return 3
    if _RE_L4.match(line):
        return 4
    return 0


# ── Main ──────────────────────────────────────────────────────────────

input_path = "/sandbox/input.md"
if not os.path.exists(input_path):
    print(json.dumps({"ok": False, "error": "/sandbox/input.md not found"}))
    raise SystemExit(1)

with open(input_path, "r", encoding="utf-8") as f:
    markdown_content = f.read()

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

lines = markdown_content.strip().split('\n')
title_done = False
i = 0

while i < len(lines):
    raw_line = lines[i]
    line = raw_line.strip()
    if not line:
        i += 1
        continue

    # Skip code blocks, tables, horizontal rules
    if line.startswith('```'):
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            i += 1
        i += 1
        continue
    if line.startswith('|') and line.endswith('|'):
        i += 1
        while i < len(lines) and lines[i].strip().startswith('|'):
            i += 1
        continue
    if re.match(r'^[-*_]{3,}$', line):
        i += 1
        continue

    # Strip blockquote marker
    line = re.sub(r'^>\s*', '', line)

    # Markdown heading patterns
    h1 = re.match(r'^#\s+(.+)', line)
    h2 = re.match(r'^##\s+(.+)', line)
    h3 = re.match(r'^###\s+(.+)', line)
    h4 = re.match(r'^####\s+(.+)', line)
    h5 = re.match(r'^#####\s+(.+)', line)

    if h1 and not title_done:
        # Document title: 方正小标宋简体, 22pt, centered
        _add_paragraph(doc, h1.group(1), '方正小标宋简体',
                       size_pt=22, bold=False,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       spacing_after=0)
        title_done = True
        _add_paragraph(doc, '', '仿宋_GB2312', size_pt=16)
    elif h2:
        # 一类标题: 黑体, 16pt, 两端对齐, 左缩进2字符, 段前0
        _add_paragraph(doc, h2.group(1), '黑体',
                       size_pt=16, bold=False,
                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       left_indent_pt=32,
                       spacing_before=0, spacing_after=0)
    elif h3:
        # 二类标题: 楷体_GB2312, 16pt, 两端对齐, 左缩进2字符, 段前0
        _add_paragraph(doc, h3.group(1), '楷体_GB2312',
                       size_pt=16, bold=False,
                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       left_indent_pt=32,
                       spacing_before=0, spacing_after=0)
    elif h4:
        # 三类标题: 仿宋_GB2312, 16pt, 两端对齐, 左缩进2字符, 段前0
        _add_paragraph(doc, h4.group(1), '仿宋_GB2312',
                       size_pt=16, bold=False,
                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       left_indent_pt=32,
                       spacing_before=0, spacing_after=0)
    elif h5:
        # Level 4 heading: 仿宋_GB2312, 16pt
        _add_paragraph(doc, h5.group(1), '仿宋_GB2312',
                       size_pt=16, bold=False,
                       spacing_before=6, spacing_after=0)
    else:
        # Content-numbered headings or body text
        level = _classify_line(line)
        if level == 1:
            # 一类标题(编号): 黑体, 16pt, 两端对齐, 左缩进2字符, 段前0
            _add_paragraph(doc, line, '黑体',
                           size_pt=16, bold=False,
                           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           left_indent_pt=32,
                           spacing_before=0, spacing_after=0)
        elif level == 2:
            # 二类标题(编号): 楷体_GB2312, 16pt, 两端对齐, 左缩进2字符, 段前0
            _add_paragraph(doc, line, '楷体_GB2312',
                           size_pt=16, bold=False,
                           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           left_indent_pt=32,
                           spacing_before=0, spacing_after=0)
        elif level in (3, 4):
            # 三类/四类标题(编号): 仿宋_GB2312, 16pt, 两端对齐, 左缩进2字符, 段前0
            _add_paragraph(doc, line, '仿宋_GB2312',
                           size_pt=16, bold=False,
                           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           left_indent_pt=32,
                           spacing_before=0, spacing_after=0)
        else:
            bullet = re.match(r'^[-*+]\s+(.+)', line)
            if bullet:
                line = '\u2022 ' + bullet.group(1)
            _add_paragraph(doc, line, '仿宋_GB2312',
                           size_pt=16, bold=False,
                           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           first_line_indent=True,
                           spacing_after=0)
    i += 1

# Write output (directory mounted writable by sandbox host)
output_path = "/sandbox/output/result.docx"
doc.save(output_path)

download_id = uuid.uuid4().hex[:16]
filename = "公文格式.docx"
print(json.dumps({"ok": True, "download_id": download_id, "filename": filename}))
